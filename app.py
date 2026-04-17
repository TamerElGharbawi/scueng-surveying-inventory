# ================================================================
#  نظام جرد معمل المساحة | Surveying Lab Inventory Agent
#  جامعة قناة السويس – كلية الهندسة
#  Version 3.0 — Auto Image Correction | Robust PDF | Fixed API
# ================================================================

import streamlit as st
import google.generativeai as genai
from PIL import Image, ImageEnhance, ImageOps, ExifTags
import io, base64, json, os, uuid, re, requests, datetime
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, Image as RLImage, PageBreak, HRFlowable,
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER
import arabic_reshaper
from bidi.algorithm import get_display

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ================================================================
# ⚙️  CONSTANTS
# ================================================================
PAGE_W, PAGE_H = A4
FONT_DIR = Path("/tmp/ar_fonts")
FONT_DIR.mkdir(exist_ok=True)
FONT_REG  = str(FONT_DIR / "Cairo-Regular.ttf")
FONT_BOLD = str(FONT_DIR / "Cairo-Bold.ttf")

UNIV_NAME  = "جامعة قناة السويس"
FAC_NAME   = "كلية الهندسة"
LAB_NAME   = "معمل المساحة"
RPT_TITLE  = "كشف جرد أجهزة ومعدات المعمل"
STATUS_OPT = ["ممتاز", "جيد جداً", "جيد", "يحتاج صيانة", "معطل"]

C_PRI  = colors.HexColor("#1a5276")
C_SEC  = colors.HexColor("#2e86c1")
C_ACC  = colors.HexColor("#f39c12")
C_LITE = colors.HexColor("#eaf0fb")
C_GRAY = colors.HexColor("#aab7b8")
C_OK   = colors.HexColor("#1e8449")
C_WARN = colors.HexColor("#e67e22")
C_ERR  = colors.HexColor("#c0392b")
C_BLK  = colors.HexColor("#1c2833")

FONT_URLS = {
    FONT_REG: [
        "https://github.com/google/fonts/raw/main/ofl/cairo/static/Cairo-Regular.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/cairo/static/Cairo-Regular.ttf",
    ],
    FONT_BOLD: [
        "https://github.com/google/fonts/raw/main/ofl/cairo/static/Cairo-Bold.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/cairo/static/Cairo-Bold.ttf",
    ],
}


# ================================================================
# 🔤  FONT SETUP
# ================================================================
@st.cache_resource
def setup_fonts():
    for path, urls in FONT_URLS.items():
        if os.path.exists(path) and os.path.getsize(path) > 10000:
            continue
        downloaded = False
        for url in urls:
            try:
                r = requests.get(url, timeout=30,
                                  headers={"User-Agent": "Mozilla/5.0"})
                r.raise_for_status()
                if len(r.content) < 10000:
                    continue
                with open(path, "wb") as f:
                    f.write(r.content)
                downloaded = True
                break
            except Exception:
                continue
        if not downloaded:
            return False, f"فشل تحميل: {Path(path).name}"
    try:
        if "Cairo" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("Cairo",      FONT_REG))
            pdfmetrics.registerFont(TTFont("Cairo-Bold", FONT_BOLD))
        return True, "ok"
    except Exception as e:
        return False, str(e)


def ar(txt):
    try:
        return get_display(arabic_reshaper.reshape(str(txt)))
    except Exception:
        return str(txt)


# ================================================================
# 🎨  PAGE CONFIG & CSS
# ================================================================
st.set_page_config(
    page_title="جرد معمل المساحة | SCU",
    page_icon="🔭",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap');
html,body,.stApp{direction:rtl!important;font-family:'Cairo',sans-serif!important;}
.main .block-container{padding:1rem .75rem!important;max-width:100%!important;}
h1,h2,h3,h4{font-family:'Cairo',sans-serif!important;text-align:right!important;color:#1a5276!important;}
.stButton>button{width:100%!important;min-height:3rem!important;font-size:1rem!important;
  font-family:'Cairo',sans-serif!important;font-weight:700!important;border-radius:12px!important;transition:all .2s!important;}
.stButton>button:hover{transform:translateY(-2px)!important;box-shadow:0 4px 14px rgba(0,0,0,.18)!important;}
.stTextInput input,.stTextArea textarea{direction:rtl!important;text-align:right!important;
  font-family:'Cairo',sans-serif!important;font-size:1rem!important;border-radius:10px!important;}
label,.stSelectbox label,.stTextArea label{font-family:'Cairo',sans-serif!important;
  font-size:1rem!important;font-weight:600!important;color:#1a5276!important;text-align:right!important;}
[data-testid="metric-container"]{direction:rtl!important;text-align:right!important;
  background:#eaf0fb!important;border-radius:12px!important;padding:.75rem!important;border-right:4px solid #1a5276!important;}
[data-testid="metric-container"] label,[data-testid="metric-container"] [data-testid="metric-value"]{
  font-family:'Cairo',sans-serif!important;color:#1a5276!important;}
.stExpander{border:1.5px solid #2e86c1!important;border-radius:12px!important;direction:rtl!important;}
.stExpander summary{font-family:'Cairo',sans-serif!important;font-weight:700!important;}
.stAlert{direction:rtl!important;font-family:'Cairo',sans-serif!important;border-radius:10px!important;}
.stTabs [data-baseweb="tab"]{font-family:'Cairo',sans-serif!important;font-size:1rem!important;font-weight:600!important;}
.stTabs [data-baseweb="tab-list"]{direction:rtl!important;}
#MainMenu,footer,header{visibility:hidden;}
::-webkit-scrollbar{width:5px;}::-webkit-scrollbar-thumb{background:#2e86c1;border-radius:3px;}
.app-header{background:linear-gradient(135deg,#1a5276,#2e86c1);color:white;
  padding:1rem 1.5rem;border-radius:14px;text-align:center;margin-bottom:1.2rem;}
.app-header h1{color:white!important;font-size:1.4rem!important;margin:0!important;}
.app-header p{color:#aed6f1!important;margin:.2rem 0 0!important;font-size:.85rem!important;}
.step-label{background:#eaf0fb;border-right:4px solid #1a5276;padding:.5rem 1rem;
  border-radius:0 8px 8px 0;font-family:'Cairo',sans-serif;font-weight:700;color:#1a5276;margin-bottom:.75rem;}
.photo-card{border:2px solid #2e86c1;border-radius:12px;padding:8px;text-align:center;background:white;direction:rtl;}
.badge-primary{background:#1a5276;color:white;padding:2px 10px;border-radius:20px;font-size:.72rem;font-family:'Cairo',sans-serif;font-weight:700;}
.badge-ref{background:#f39c12;color:white;padding:2px 10px;border-radius:20px;font-size:.72rem;font-family:'Cairo',sans-serif;font-weight:700;}
.dup-warning{background:#fdedec;border:2px solid #c0392b;border-radius:12px;padding:1rem;direction:rtl;animation:dupPulse 1.4s infinite;}
@keyframes dupPulse{0%{box-shadow:0 0 0 0 rgba(192,57,43,.5)}70%{box-shadow:0 0 0 10px rgba(192,57,43,0)}100%{box-shadow:0 0 0 0 rgba(192,57,43,0)}}
.success-flash{background:linear-gradient(135deg,#1e8449,#27ae60);color:white;padding:1rem;border-radius:12px;text-align:center;font-family:'Cairo',sans-serif;font-weight:700;font-size:1.05rem;}
.key-guide{background:#fef9e7;border:2px dashed #f39c12;border-radius:12px;padding:1rem;direction:rtl;font-family:'Cairo',sans-serif;line-height:2;}
.editor-box{background:#f8f9fa;border:1.5px solid #2e86c1;border-radius:14px;padding:1rem;margin:8px 0;}
</style>
""", unsafe_allow_html=True)


# ================================================================
# 🗃️  SESSION STATE
# ================================================================
def init_state():
    defs = dict(
        inventory=[], captured_photos=[], photo_hashes=set(),
        page="capture", professor_name="",
        report_date=datetime.date.today().strftime("%Y/%m/%d"),
        ai_result=None, gemini_key="", edit_photo_idx=None,
    )
    for k, v in defs.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


# ================================================================
# 🔑  API KEY VALIDATION
# ================================================================
def validate_key(key):
    key = (key or "").strip()
    if not key:
        return False, "المفتاح فارغ"
    if key.startswith("AQ.") or key.startswith("ya29."):
        return False, (
            "⛔ هذا مفتاح OAuth — لا يعمل مع Gemini API.\n"
            "✅ احصل على مفتاح صحيح من: aistudio.google.com/app/apikey\n"
            "   المفتاح الصحيح يبدأ دائماً بـ AIzaSy"
        )
    if not key.startswith("AIza"):
        return False, (
            "⚠️ صيغة المفتاح غير متوقعة — يجب أن يبدأ بـ AIzaSy\n"
            "تأكد من النسخ الكامل من aistudio.google.com/app/apikey"
        )
    return True, "ok"


# ================================================================
# 🖼️  IMAGE PROCESSING
# ================================================================
def fix_exif_rotation(img):
    try:
        exif = img._getexif()
        if exif is None:
            return img
        orient_tag = next((k for k,v in ExifTags.TAGS.items() if v=="Orientation"), None)
        if orient_tag and orient_tag in exif:
            o = exif[orient_tag]
            if   o == 3: img = img.rotate(180, expand=True)
            elif o == 6: img = img.rotate(270, expand=True)
            elif o == 8: img = img.rotate(90,  expand=True)
            elif o == 2: img = img.transpose(Image.FLIP_LEFT_RIGHT)
            elif o == 4: img = img.transpose(Image.FLIP_TOP_BOTTOM)
            elif o == 5: img = img.transpose(Image.FLIP_LEFT_RIGHT).rotate(90,  expand=True)
            elif o == 7: img = img.transpose(Image.FLIP_LEFT_RIGHT).rotate(270, expand=True)
    except Exception:
        pass
    return img


def ai_detect_rotation(img_bytes, api_key):
    """Ask Gemini for correct rotation in degrees CW."""
    try:
        genai.configure(api_key=api_key.strip())
        model = genai.GenerativeModel("gemini-1.5-flash")
        pil = Image.open(io.BytesIO(img_bytes))
        if max(pil.size) > 900:
            pil.thumbnail((900, 900), Image.LANCZOS)
        buf = io.BytesIO(); pil.save(buf, "JPEG", quality=75)
        img_part = {"mime_type":"image/jpeg",
                    "data": base64.b64encode(buf.getvalue()).decode()}
        resp = model.generate_content([
            "How many degrees clockwise should this image be rotated to appear correctly upright? "
            "The device/equipment/text should read normally. "
            "Reply with ONLY one integer: 0, 90, 180, or 270.",
            img_part
        ])
        m = re.search(r"\b(0|90|180|270)\b", resp.text)
        return int(m.group()) if m else 0
    except Exception:
        return 0


def ai_smart_crop(img_bytes, api_key):
    """Ask Gemini for tight crop bounding box. Returns (left%, top%, right%, bottom%) as 0-1."""
    try:
        genai.configure(api_key=api_key.strip())
        model = genai.GenerativeModel("gemini-1.5-flash")
        pil = Image.open(io.BytesIO(img_bytes))
        if max(pil.size) > 900:
            pil.thumbnail((900, 900), Image.LANCZOS)
        buf = io.BytesIO(); pil.save(buf, "JPEG", quality=75)
        img_part = {"mime_type":"image/jpeg",
                    "data": base64.b64encode(buf.getvalue()).decode()}
        resp = model.generate_content([
            "Identify the main subject (device/equipment/nameplate) in this image. "
            "Return ONLY a JSON with crop percentages (0-100): "
            '{"left":5,"top":5,"right":95,"bottom":95} '
            "Add ~5% padding. No backticks, no extra text.",
            img_part
        ])
        raw = re.sub(r"```json|```", "", resp.text).strip()
        d = json.loads(raw)
        l,t,r,b = d["left"],d["top"],d["right"],d["bottom"]
        if 0 <= l < r <= 100 and 0 <= t < b <= 100:
            return l/100, t/100, r/100, b/100
    except Exception:
        pass
    return 0.0, 0.0, 1.0, 1.0


def auto_enhance(img):
    img = ImageOps.autocontrast(img, cutoff=1)
    img = ImageEnhance.Sharpness(img).enhance(1.4)
    img = ImageEnhance.Contrast(img).enhance(1.12)
    return img


def process_image(data, rotate_deg=0, crop_box=(0.0,0.0,1.0,1.0), enhance=True):
    img = Image.open(io.BytesIO(data))
    if img.mode not in ("RGB","L"):
        img = img.convert("RGB")
    img = fix_exif_rotation(img)
    if rotate_deg:
        img = img.rotate(-rotate_deg, expand=True)  # negative = clockwise
    l,t,r,b = crop_box
    if (l,t,r,b) != (0.0,0.0,1.0,1.0):
        W,H = img.size
        img = img.crop((int(l*W), int(t*H), int(r*W), int(b*H)))
    if enhance:
        img = auto_enhance(img)
    if max(img.size) > 1400:
        img.thumbnail((1400,1400), Image.LANCZOS)
    buf = io.BytesIO(); img.save(buf, "JPEG", quality=88)
    return buf.getvalue()


def resize_img(data, mx=1200):
    img = Image.open(io.BytesIO(data))
    if img.mode not in ("RGB","L"):
        img = img.convert("RGB")
    img = fix_exif_rotation(img)
    if max(img.size) > mx:
        img.thumbnail((mx,mx), Image.LANCZOS)
    buf = io.BytesIO(); img.save(buf, "JPEG", quality=85)
    return buf.getvalue()


def thumb(data, sz=(160,160)):
    img = Image.open(io.BytesIO(data))
    if img.mode not in ("RGB","L"):
        img = img.convert("RGB")
    img.thumbnail(sz, Image.LANCZOS)
    canvas = Image.new("RGB", sz, (255,255,255))
    canvas.paste(img, ((sz[0]-img.size[0])//2, (sz[1]-img.size[1])//2))
    buf = io.BytesIO(); canvas.save(buf, "JPEG", quality=85)
    return buf.getvalue()


def b64img(data):
    return base64.b64encode(data).decode()


# ================================================================
# 🤖  GEMINI: DEVICE DATA EXTRACTION
# ================================================================
def extract_device_info(images, api_key):
    valid, msg = validate_key(api_key)
    if not valid:
        return {"error": msg}
    try:
        genai.configure(api_key=api_key.strip())
        model = genai.GenerativeModel("gemini-1.5-flash")
        parts = ["""أنت نظام متخصص في التعرف على الأجهزة والمعدات الهندسية.
حلّل الصور واستخرج البيانات. أجب بـ JSON صالح فقط بلا نص خارجي:
{"device_type":"نوع الجهاز","brand":"الماركة","serial_number":"الرقم التسلسلي أو غير مقروء",
"condition":"ممتاز أو جيد جداً أو جيد أو يحتاج صيانة أو معطل","notes":"ملاحظات","confidence":"high أو medium أو low"}"""]
        for img_bytes in images:
            pil = Image.open(io.BytesIO(img_bytes))
            if pil.mode not in ("RGB","L"): pil = pil.convert("RGB")
            if max(pil.size) > 1200: pil.thumbnail((1200,1200), Image.LANCZOS)
            buf = io.BytesIO(); pil.save(buf, "JPEG", quality=85)
            parts.append({"mime_type":"image/jpeg",
                           "data": base64.b64encode(buf.getvalue()).decode()})
        resp = model.generate_content(parts)
        raw = re.sub(r"```json|```","",resp.text).strip()
        return json.loads(raw)
    except json.JSONDecodeError as e:
        return {"device_type":"","brand":"","serial_number":"",
                "condition":"جيد","notes":f"تعذّر تحليل الاستجابة","confidence":"low"}
    except Exception as e:
        err = str(e)
        if "API_KEY_INVALID" in err or "400" in err:
            return {"error":"❌ مفتاح API غير صالح — تأكد من نسخه من aistudio.google.com/app/apikey"}
        if "quota" in err.lower() or "429" in err:
            return {"error":"⏱️ تجاوزت الحد المجاني — انتظر دقيقة وأعد المحاولة"}
        return {"error": f"خطأ Gemini: {err}"}


# ================================================================
# 🔍  DUPLICATE DETECTION
# ================================================================
def is_duplicate(serial, exclude_idx=-1):
    if not serial or serial.strip().upper() in ["","غير مقروء","غير محدد"]:
        return False
    for i, item in enumerate(st.session_state.inventory):
        if i == exclude_idx: continue
        if item.get("serial_number","").strip().upper() == serial.strip().upper():
            return True
    return False


def get_duplicates():
    seen = {}
    for i, item in enumerate(st.session_state.inventory):
        sn = item.get("serial_number","").strip().upper()
        if sn and sn not in ["غير مقروء","غير محدد"]:
            seen.setdefault(sn, []).append(i+1)
    return [(sn,idx) for sn,idx in seen.items() if len(idx)>1]


# ================================================================
# 📄  PDF REPORT
# ================================================================
class PDFReport:
    def __init__(self, inv, prof, date):
        self.inv=inv; self.prof=prof; self.date=date
        self.dups=get_duplicates(); self.buf=io.BytesIO()

    def _ps(self, size=10, bold=False, align=TA_RIGHT, color=C_BLK):
        return ParagraphStyle(f"s{uuid.uuid4().hex[:4]}",
            fontName="Cairo-Bold" if bold else "Cairo",
            fontSize=size, textColor=color, alignment=align,
            leading=int(size*1.7), wordWrap="CJK")

    def _p(self, txt, **kw):
        return Paragraph(ar(str(txt)), self._ps(**kw))

    def _hf(self, canvas, doc):
        canvas.saveState()
        W,H = A4
        canvas.setFillColor(C_PRI)
        canvas.rect(0, H-2.2*cm, W, 2.2*cm, fill=1, stroke=0)
        canvas.setFillColor(colors.white); canvas.setStrokeColor(C_ACC); canvas.setLineWidth(1.2)
        canvas.rect(W-2.8*cm, H-2.0*cm, 2.2*cm, 1.8*cm, fill=1, stroke=1)
        canvas.setFillColor(C_GRAY); canvas.setFont("Cairo",6)
        canvas.drawCentredString(W-1.7*cm, H-1.2*cm, ar("شعار"))
        canvas.drawCentredString(W-1.7*cm, H-1.5*cm, ar("الجامعة"))
        canvas.setFillColor(colors.white); canvas.setFont("Cairo-Bold",10.5)
        canvas.drawRightString(W-3.2*cm, H-0.95*cm, ar(f"{UNIV_NAME}  |  {FAC_NAME}"))
        canvas.setFont("Cairo",8.5)
        canvas.drawRightString(W-3.2*cm, H-1.55*cm, ar(f"{LAB_NAME}   |   {self.date}"))
        if self.prof:
            canvas.setFillColor(colors.HexColor("#aed6f1")); canvas.setFont("Cairo",8)
            canvas.drawString(0.6*cm, H-1.2*cm, ar(f"أ.د / {self.prof}"))
        canvas.setFillColor(C_PRI)
        canvas.rect(0, 0, W, 1.1*cm, fill=1, stroke=0)
        canvas.setFillColor(colors.white); canvas.setFont("Cairo",7.5)
        canvas.drawCentredString(W/2, 0.38*cm, ar(f"{UNIV_NAME} — {FAC_NAME} — {LAB_NAME}"))
        canvas.setFont("Cairo-Bold",8)
        canvas.drawString(0.7*cm, 0.38*cm, ar(f"صفحة {doc.page}"))
        canvas.restoreState()

    def _cover(self):
        el=[Spacer(1,2.2*cm)]
        lbox=Table([[self._p("[ شعار جامعة قناة السويس ]",size=10,align=TA_CENTER,color=C_GRAY)]],
                   colWidths=[6*cm],rowHeights=[5*cm])
        lbox.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                                   ("BOX",(0,0),(-1,-1),2,C_ACC),("BACKGROUND",(0,0),(-1,-1),C_LITE)]))
        wrap=Table([[lbox]],colWidths=[PAGE_W-4*cm])
        wrap.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"CENTER")]))
        el+=[wrap,Spacer(1,.7*cm)]
        el.append(self._p(UNIV_NAME,size=19,bold=True,align=TA_CENTER,color=C_PRI))
        el+=[Spacer(1,.2*cm),self._p(FAC_NAME,size=14,bold=True,align=TA_CENTER,color=C_SEC),
             Spacer(1,.15*cm),self._p(LAB_NAME,size=12,align=TA_CENTER),Spacer(1,.9*cm),
             HRFlowable(width="70%",thickness=2,color=C_ACC,hAlign="CENTER"),Spacer(1,.7*cm),
             self._p(RPT_TITLE,size=16,bold=True,align=TA_CENTER),Spacer(1,.35*cm),
             self._p(f"العام الدراسي {datetime.datetime.now().year-1} / {datetime.datetime.now().year}",
                     size=11,align=TA_CENTER),Spacer(1,1.6*cm)]
        rows=[[self._p("تاريخ الجرد :",size=11,bold=True),self._p(self.date,size=11)],
               [self._p("المُعِد :",size=11,bold=True),self._p(self.prof or "—",size=11)],
               [self._p("إجمالي الأجهزة :",size=11,bold=True),
                self._p(str(len(self.inv)),size=12,bold=True,color=C_PRI)]]
        info=Table(rows,colWidths=[4.5*cm,8*cm])
        info.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"RIGHT"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("ROWBACKGROUNDS",(0,0),(-1,-1),[C_LITE,colors.white,C_LITE]),
            ("BOX",(0,0),(-1,-1),1,C_SEC),("INNERGRID",(0,0),(-1,-1),.5,C_GRAY),
            ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("RIGHTPADDING",(0,0),(-1,-1),8),("LEFTPADDING",(0,0),(-1,-1),8)]))
        outer=Table([[info]],colWidths=[PAGE_W-4*cm])
        outer.setStyle(TableStyle([("ALIGN",(0,0),(-1,-1),"CENTER")]))
        el.append(outer); return el

    def _stats(self):
        el=[self._p("الملخص الإحصائي",size=13,bold=True,color=C_PRI),
             HRFlowable(width="100%",thickness=1,color=C_SEC),Spacer(1,.3*cm)]
        brands={}; conds={}
        for item in self.inv:
            brands[item.get("brand","غير محدد")]=brands.get(item.get("brand","غير محدد"),0)+1
            conds[item.get("condition","غير محدد")]=conds.get(item.get("condition","غير محدد"),0)+1
        need=conds.get("يحتاج صيانة",0)+conds.get("معطل",0)
        excel=conds.get("ممتاز",0)+conds.get("جيد جداً",0)
        rows=[[self._p("البيان",size=10,bold=True,color=colors.white),
                self._p("القيمة",size=10,bold=True,color=colors.white)],
               [self._p("إجمالي عدد الأجهزة",size=10),self._p(str(len(self.inv)),size=10,bold=True)],
               [self._p("عدد الماركات",size=10),self._p(str(len(brands)),size=10,bold=True)],
               [self._p("ممتاز / جيد جداً",size=10),self._p(str(excel),size=10,bold=True,color=C_OK)],
               [self._p("تحتاج صيانة / معطلة",size=10),self._p(str(need),size=10,bold=True,color=C_ERR)]]
        if self.dups:
            rows.append([self._p("أرقام تسلسلية مكررة",size=10,color=C_ERR),
                          self._p(str(len(self.dups)),size=10,bold=True,color=C_ERR)])
        tbl=Table(rows,colWidths=[11*cm,4*cm])
        tbl.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),C_PRI),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_LITE,colors.white]),
            ("BOX",(0,0),(-1,-1),1,C_PRI),("INNERGRID",(0,0),(-1,-1),.5,C_GRAY),
            ("ALIGN",(0,0),(-1,-1),"RIGHT"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
            ("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6)]))
        el.append(tbl)
        if brands:
            el+=[Spacer(1,.35*cm),
                  self._p("توزيع الماركات:  "+"  |  ".join(
                    f"{b}: {c}" for b,c in sorted(brands.items(),key=lambda x:-x[1])),
                    size=9,color=C_SEC)]
        return el

    def _inv_table(self):
        el=[PageBreak(),self._p("كشف الأجهزة والمعدات",size=13,bold=True,color=C_PRI),
             HRFlowable(width="100%",thickness=1,color=C_SEC),Spacer(1,.3*cm)]
        CW=[1*cm,3.3*cm,2.5*cm,2.4*cm,2*cm,3*cm,3*cm]
        hdr=[self._p("م",size=9,bold=True,align=TA_CENTER,color=colors.white),
              self._p("نوع الجهاز",size=9,bold=True,color=colors.white),
              self._p("الماركة / المصنّع",size=9,bold=True,color=colors.white),
              self._p("الرقم التسلسلي",size=9,bold=True,color=colors.white),
              self._p("الحالة",size=9,bold=True,align=TA_CENTER,color=colors.white),
              self._p("الملاحظات",size=9,bold=True,color=colors.white),
              self._p("صورة الجهاز",size=9,bold=True,align=TA_CENTER,color=colors.white)]
        data=[hdr]; ROW_H=3.2*cm
        for i,item in enumerate(self.inv):
            cond=item.get("condition","")
            cc=(C_OK if cond in["ممتاز","جيد جداً"] else
                C_SEC if cond=="جيد" else
                C_WARN if cond=="يحتاج صيانة" else C_ERR)
            ph=None
            for p in item.get("photos",[]):
                if p.get("is_primary"): ph=p["data"]; break
            if ph is None and item.get("photos"): ph=item["photos"][0]["data"]
            if ph:
                t=thumb(ph,(113,113))
                img_cell=RLImage(io.BytesIO(t),width=3*cm,height=3*cm)
            else:
                img_cell=self._p("لا توجد صورة",size=8,align=TA_CENTER,color=C_GRAY)
            data.append([
                self._p(str(i+1),size=9,align=TA_CENTER),
                self._p(item.get("device_type","—"),size=9),
                self._p(item.get("brand","—"),size=9),
                self._p(item.get("serial_number","—"),size=8),
                self._p(cond,size=8,align=TA_CENTER,color=cc),
                self._p(item.get("notes","—"),size=8),
                img_cell])
        tbl=Table(data,colWidths=CW,rowHeights=[0.7*cm]+[ROW_H]*len(self.inv),repeatRows=1)
        tstyle=[("BACKGROUND",(0,0),(-1,0),C_PRI),("BOX",(0,0),(-1,-1),1.5,C_PRI),
                ("INNERGRID",(0,0),(-1,-1),.5,C_GRAY),
                ("ALIGN",(0,0),(0,-1),"CENTER"),("ALIGN",(4,1),(4,-1),"CENTER"),
                ("ALIGN",(6,1),(6,-1),"CENTER"),("ALIGN",(0,0),(-1,0),"RIGHT"),
                ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
                ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4)]
        for r in range(1,len(data)):
            tstyle.append(("BACKGROUND",(0,r),(-1,r),C_LITE if r%2==0 else colors.white))
        tbl.setStyle(TableStyle(tstyle)); el.append(tbl); return el

    def _dups_table(self):
        el=[PageBreak(),self._p("تحذير: أرقام تسلسلية مكررة",size=13,bold=True,color=C_ERR),
             HRFlowable(width="100%",thickness=2,color=C_ERR),Spacer(1,.3*cm),
             self._p("يرجى مراجعة الأجهزة التالية:",size=10),Spacer(1,.3*cm)]
        rows=[[self._p("الرقم التسلسلي",size=10,bold=True,color=colors.white),
                self._p("أرقام الصفوف",size=10,bold=True,color=colors.white)]]
        for sn,idx in self.dups:
            rows.append([self._p(sn,size=10,color=C_ERR),
                          self._p("، ".join(str(x) for x in idx),size=10)])
        tbl=Table(rows,colWidths=[9*cm,6*cm])
        tbl.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),C_ERR),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#fdedec"),colors.white]),
            ("BOX",(0,0),(-1,-1),1.5,C_ERR),("INNERGRID",(0,0),(-1,-1),.5,C_GRAY),
            ("ALIGN",(0,0),(-1,-1),"RIGHT"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7)]))
        el.append(tbl); return el

    def _sig(self):
        el=[Spacer(1,2*cm),HRFlowable(width="100%",thickness=.5,color=C_GRAY),Spacer(1,.5*cm)]
        rows=[[self._p("اعتُمد بمعرفة:",size=10),self._p("المراجع / الرئيس المباشر:",size=10)],
               [self._p("أ.د / "+(self.prof or "............................"),size=11,bold=True),
                self._p("د / ....................................",size=11)],
               [Spacer(1,1.4*cm),Spacer(1,1.4*cm)],
               [self._p("التوقيع:  ___________________",size=10,color=C_GRAY),
                self._p("التوقيع:  ___________________",size=10,color=C_GRAY)],
               [self._p(f"التاريخ:  {self.date}",size=10),
                self._p("التاريخ:  _____  /  _____  /  _____",size=10)]]
        tbl=Table(rows,colWidths=[8.5*cm,8.5*cm])
        tbl.setStyle(TableStyle([("ALIGN",(0,0),(0,-1),"RIGHT"),("ALIGN",(1,0),(1,-1),"LEFT"),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),("TOPPADDING",(0,0),(-1,-1),6),
            ("BOTTOMPADDING",(0,0),(-1,-1),4)]))
        el+=[tbl,Spacer(1,.4*cm),
              self._p(f"هذا الكشف صادر من إدارة {LAB_NAME} — {FAC_NAME} — {UNIV_NAME}",
                      size=8,align=TA_CENTER,color=C_GRAY)]
        return el

    def build(self):
        doc=SimpleDocTemplate(self.buf,pagesize=A4,
            rightMargin=2*cm,leftMargin=2*cm,topMargin=2.7*cm,bottomMargin=1.8*cm)
        story=self._cover()+[PageBreak()]+self._stats()+self._inv_table()
        if self.dups: story+=self._dups_table()
        story+=self._sig()
        doc.build(story,onFirstPage=self._hf,onLaterPages=self._hf)
        return self.buf.getvalue()


# ================================================================
# 📝  DOCX REPORT
# ================================================================
class DOCXReport:
    def __init__(self,inv,prof,date):
        self.inv=inv; self.prof=prof; self.date=date
        self.dups=get_duplicates(); self.doc=Document()

    def _rtl(self,para):
        pPr=para._p.get_or_add_pPr()
        b=OxmlElement("w:bidi"); b.set(qn("w:val"),"1"); pPr.append(b)
        para.alignment=WD_ALIGN_PARAGRAPH.RIGHT

    def _run(self,para,text,bold=False,size=11,color=None):
        run=para.add_run(text); run.bold=bold
        run.font.name="Cairo"; run.font.size=Pt(size)
        run._element.rPr.rFonts.set(qn("w:hint"),"cs")
        if color: run.font.color.rgb=RGBColor.from_string(color)
        return run

    def _para(self,text,size=11,bold=False,color=None):
        p=self.doc.add_paragraph(); self._rtl(p)
        self._run(p,text,bold=bold,size=size,color=color); return p

    def _heading(self,text,size=14,color="1a5276"):
        p=self.doc.add_paragraph(); self._rtl(p)
        self._run(p,text,bold=True,size=size,color=color); return p

    def _shd(self,cell,fill_hex):
        shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill_hex); shd.set(qn("w:val"),"clear")
        cell._tc.get_or_add_tcPr().append(shd)

    def build(self):
        sec=self.doc.sections[0]
        sec.right_margin=Cm(2); sec.left_margin=Cm(2)
        sec.top_margin=Cm(3); sec.bottom_margin=Cm(2)
        self.doc.styles["Normal"].font.name="Cairo"
        self.doc.styles["Normal"].font.size=Pt(11)

        # header
        hp=sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        self._rtl(hp)
        hr=hp.add_run(f"{UNIV_NAME}  |  {FAC_NAME}  |  {LAB_NAME}  |  {self.date}")
        hr.font.name="Cairo"; hr.font.size=Pt(9); hr.bold=True
        hr.font.color.rgb=RGBColor(0x1a,0x52,0x76)
        hr._element.rPr.rFonts.set(qn("w:hint"),"cs")
        fp=sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
        fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        fr=fp.add_run(f"{UNIV_NAME} — {FAC_NAME} — {LAB_NAME}")
        fr.font.name="Cairo"; fr.font.size=Pt(8)
        fr.font.color.rgb=RGBColor(0x1a,0x52,0x76)
        fr._element.rPr.rFonts.set(qn("w:hint"),"cs")

        # cover
        cp=self.doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cr=cp.add_run("[ شعار جامعة قناة السويس ]")
        cr.font.size=Pt(14); cr.font.color.rgb=RGBColor(0xaa,0xbb,0xcc)
        cr._element.rPr.rFonts.set(qn("w:hint"),"cs")
        self.doc.add_paragraph()
        self._heading(UNIV_NAME,22,"1a5276"); self._heading(FAC_NAME,16,"2e86c1")
        self._heading(LAB_NAME,13,"1c2833"); self.doc.add_paragraph()
        d=self.doc.add_paragraph("─"*42); d.alignment=WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()
        self._heading(RPT_TITLE,16,"1c2833"); self.doc.add_paragraph()
        ct=self.doc.add_table(3,2); ct.style="Table Grid"
        ct.alignment=WD_TABLE_ALIGNMENT.CENTER
        for r,(lbl,val) in enumerate([("تاريخ الجرد:",self.date),
            ("المُعِد:",self.prof or "—"),("إجمالي الأجهزة:",str(len(self.inv)))]):
            for c,txt in enumerate([lbl,val]):
                p=ct.rows[r].cells[c].paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rtl(p)
                run=p.add_run(txt); run.bold=(c==0)
                run.font.name="Cairo"; run.font.size=Pt(11)
                run._element.rPr.rFonts.set(qn("w:hint"),"cs")

        self.doc.add_page_break()

        # stats
        self._heading("الملخص الإحصائي",13,"1a5276")
        brands={}; conds={}
        for item in self.inv:
            brands[item.get("brand","غير محدد")]=brands.get(item.get("brand","غير محدد"),0)+1
            conds[item.get("condition","غير محدد")]=conds.get(item.get("condition","غير محدد"),0)+1
        rows_data=[("إجمالي عدد الأجهزة",str(len(self.inv))),
                    ("عدد الماركات",str(len(brands))),
                    ("تحتاج صيانة / معطلة",str(conds.get("يحتاج صيانة",0)+conds.get("معطل",0)))]
        if self.dups: rows_data.append(("أرقام مكررة",str(len(self.dups))))
        st2=self.doc.add_table(len(rows_data)+1,2); st2.style="Table Grid"
        st2.alignment=WD_TABLE_ALIGNMENT.RIGHT
        for c,txt in enumerate(["البيان","القيمة"]):
            cell=st2.rows[0].cells[c]; p=cell.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rtl(p)
            run=p.add_run(txt); run.bold=True
            run.font.name="Cairo"; run.font.size=Pt(11)
            run._element.rPr.rFonts.set(qn("w:hint"),"cs")
            self._shd(cell,"1a5276")
        for r,(lbl,val) in enumerate(rows_data):
            for c,txt in enumerate([lbl,val]):
                p=st2.rows[r+1].cells[c].paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rtl(p)
                run=p.add_run(txt); run.font.name="Cairo"; run.font.size=Pt(11)
                run._element.rPr.rFonts.set(qn("w:hint"),"cs")
        self.doc.add_paragraph()
        self.doc.add_page_break()

        # inventory table
        self._heading("كشف الأجهزة والمعدات",14,"1a5276")
        hdrs=["م","نوع الجهاز","الماركة / المصنّع","الرقم التسلسلي","الحالة","الملاحظات","صورة"]
        CW=[Cm(1),Cm(3.5),Cm(2.5),Cm(2.5),Cm(2),Cm(3),Cm(2.5)]
        tbl=self.doc.add_table(len(self.inv)+1,7)
        tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.RIGHT
        for c,(h,w) in enumerate(zip(hdrs,CW)):
            cell=tbl.rows[0].cells[c]; cell.width=w
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; self._rtl(p)
            run=p.add_run(h); run.bold=True
            run.font.name="Cairo"; run.font.size=Pt(10)
            run._element.rPr.rFonts.set(qn("w:hint"),"cs")
            self._shd(cell,"1a5276")
        for r,item in enumerate(self.inv):
            row=tbl.rows[r+1]
            trPr=row._tr.get_or_add_trPr()
            trH=OxmlElement("w:trHeight")
            trH.set(qn("w:val"),"1700"); trH.set(qn("w:hRule"),"exact"); trPr.append(trH)
            vals=[str(r+1),item.get("device_type","—"),item.get("brand","—"),
                  item.get("serial_number","—"),item.get("condition","—"),
                  item.get("notes","—"),None]
            for c,val in enumerate(vals):
                cell=row.cells[c]; cell.width=CW[c]
                p=cell.paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER if c in(0,6) else WD_ALIGN_PARAGRAPH.RIGHT
                self._rtl(p)
                if c==6:
                    ph=None
                    for pi in item.get("photos",[]):
                        if pi.get("is_primary"): ph=pi["data"]; break
                    if ph is None and item.get("photos"): ph=item["photos"][0]["data"]
                    if ph:
                        try:
                            run=p.add_run()
                            run.add_picture(io.BytesIO(thumb(ph,(150,150))),
                                            width=Cm(2.4),height=Cm(2.4))
                        except Exception:
                            run=p.add_run("—"); run.font.name="Cairo"
                            run._element.rPr.rFonts.set(qn("w:hint"),"cs")
                    else:
                        run=p.add_run("—"); run.font.name="Cairo"
                        run._element.rPr.rFonts.set(qn("w:hint"),"cs")
                else:
                    run=p.add_run(str(val) if val else "—")
                    run.font.name="Cairo"; run.font.size=Pt(9)
                    run._element.rPr.rFonts.set(qn("w:hint"),"cs")
                if r%2==0: self._shd(cell,"eaf0fb")
        self.doc.add_paragraph()

        # duplicates
        if self.dups:
            self.doc.add_page_break()
            self._heading("أجهزة بأرقام تسلسلية مكررة",13,"c0392b")
            dt=self.doc.add_table(len(self.dups)+1,2); dt.style="Table Grid"
            for c,txt in enumerate(["الرقم التسلسلي","أرقام الصفوف"]):
                cell=dt.rows[0].cells[c]; p=cell.paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rtl(p)
                run=p.add_run(txt); run.bold=True
                run.font.name="Cairo"; run.font.size=Pt(11)
                run._element.rPr.rFonts.set(qn("w:hint"),"cs")
                self._shd(cell,"c0392b")
            for r,(sn,idx) in enumerate(self.dups):
                for c,txt in enumerate([sn,"، ".join(str(x) for x in idx)]):
                    p=dt.rows[r+1].cells[c].paragraphs[0]
                    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rtl(p)
                    run=p.add_run(txt); run.font.name="Cairo"; run.font.size=Pt(10)
                    run._element.rPr.rFonts.set(qn("w:hint"),"cs")
            self.doc.add_paragraph()

        # signature
        self.doc.add_paragraph()
        sig=self.doc.add_table(4,2)
        for r,(l,rr) in enumerate([
            ("اعتُمد بمعرفة:","المراجع / الرئيس المباشر:"),
            (f"أ.د / {self.prof or '......................'}","د / ......................"),
            ("",""),
            ("التوقيع: _______________","التوقيع: _______________")]):
            for c,txt in enumerate([l,rr]):
                p=sig.rows[r].cells[c].paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT if c==0 else WD_ALIGN_PARAGRAPH.LEFT
                self._rtl(p)
                run=p.add_run(txt); run.font.name="Cairo"
                run.font.size=Pt(11); run.bold=(r==1)
                run._element.rPr.rFonts.set(qn("w:hint"),"cs")
        fp=self.doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; self._rtl(fp)
        run=fp.add_run(f"هذا الكشف صادر من إدارة {LAB_NAME} — {FAC_NAME} — {UNIV_NAME}")
        run.font.name="Cairo"; run.font.size=Pt(9)
        run.font.color.rgb=RGBColor(0x99,0x99,0x99)
        run._element.rPr.rFonts.set(qn("w:hint"),"cs")

        buf=io.BytesIO(); self.doc.save(buf); return buf.getvalue()


# ================================================================
# 📱  UI HELPERS
# ================================================================
def nav():
    cols=st.columns(3)
    items=[("📷","إضافة جهاز","capture"),
           ("📋",f"القائمة ({len(st.session_state.inventory)})","list"),
           ("📤","تصدير","export")]
    for col,(icon,label,pg) in zip(cols,items):
        active="✅ " if st.session_state.page==pg else ""
        with col:
            if st.button(f"{active}{icon} {label}",key=f"nav_{pg}"):
                st.session_state.page=pg; st.rerun()


def add_photo(data,is_primary,label):
    h=hash(data)
    if h not in st.session_state.photo_hashes:
        st.session_state.photo_hashes.add(h)
        st.session_state.captured_photos.append(
            {"id":uuid.uuid4().hex[:8],"data":data,"is_primary":is_primary,"label":label})


# ================================================================
# ✏️  IMAGE EDITOR
# ================================================================
def image_editor(idx):
    ph=st.session_state.captured_photos[idx]
    data=ph["data"]
    st.markdown(f'<div class="editor-box">', unsafe_allow_html=True)
    st.markdown(f"**✏️ محرر الصورة: {ph['label']}**")

    c1,c2=st.columns([3,1])
    with c1:
        rotate=st.select_slider("↩️ تدوير (درجة عكس عقارب الساعة)",
            options=[0,90,180,270],value=0,key=f"rot_{ph['id']}")
    with c2:
        enhance_on=st.checkbox("✨ تحسين",value=True,key=f"enh_{ph['id']}")

    cc1,cc2=st.columns(2)
    cl=cc1.slider("✂️ قص يسار %", 0,40,0,key=f"cl_{ph['id']}")
    cr=cc1.slider("✂️ قص يمين %",0,40,0,key=f"cr_{ph['id']}")
    ct=cc2.slider("✂️ قص أعلى %", 0,40,0,key=f"ct_{ph['id']}")
    cb=cc2.slider("✂️ قص أسفل %",0,40,0,key=f"cb_{ph['id']}")
    crop_box=(cl/100,ct/100,1.0-cr/100,1.0-cb/100)

    # live preview
    try:
        preview=process_image(data,rotate_deg=rotate,crop_box=crop_box,enhance=enhance_on)
        st.image(preview,use_container_width=True,caption="معاينة حية")
    except Exception as e:
        st.error(f"خطأ معاينة: {e}"); preview=data

    ba,bb,bc=st.columns(3)
    with ba:
        if st.button("🤖 AI تصحيح تلقائي",key=f"aifix_{ph['id']}",
                     disabled=not st.session_state.gemini_key):
            with st.spinner("Gemini يحلل الصورة..."):
                rot=ai_detect_rotation(data,st.session_state.gemini_key)
                l,t,r,b=ai_smart_crop(data,st.session_state.gemini_key)
                fixed=process_image(data,rotate_deg=rot,crop_box=(l,t,r,b),enhance=True)
                st.session_state.captured_photos[idx]["data"]=fixed
                st.session_state.edit_photo_idx=None
            st.success(f"✅ تدوير {rot}° + قص ذكي + تحسين")
            st.rerun()
    with bb:
        if st.button("✅ تطبيق",key=f"apply_{ph['id']}",type="primary"):
            try:
                processed=process_image(data,rotate_deg=rotate,
                                         crop_box=crop_box,enhance=enhance_on)
                st.session_state.captured_photos[idx]["data"]=processed
                st.session_state.edit_photo_idx=None
                st.rerun()
            except Exception as e:
                st.error(f"خطأ: {e}")
    with bc:
        if st.button("↩️ إلغاء",key=f"cancel_{ph['id']}"):
            st.session_state.edit_photo_idx=None; st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


# ================================================================
# 📷  PAGE: CAPTURE
# ================================================================
def page_capture():

    with st.expander("⚙️ إعدادات التقرير والـ API",
                     expanded=not st.session_state.gemini_key):
        st.markdown("""
        <div class="key-guide">
          <strong>🔑 كيف تحصل على مفتاح Gemini الصحيح؟</strong><br>
          1️⃣ افتح المتصفح واذهب إلى: <code>aistudio.google.com/app/apikey</code><br>
          2️⃣ اضغط <strong>"Create API Key"</strong> ثم اختر مشروعاً<br>
          3️⃣ انسخ المفتاح — يبدأ بـ <code>AIzaSy...</code> (وليس AQ. أو ya29.)<br>
          ⚠️ مفتاح <code>AQ.Ab8...</code> هو مفتاح OAuth — <strong>لا يعمل</strong> مع Gemini API
        </div>
        """, unsafe_allow_html=True)

        key_input=st.text_input("🔑 Gemini API Key",
            value=st.session_state.gemini_key,type="password",placeholder="AIzaSy...")
        if key_input:
            valid,msg=validate_key(key_input)
            if not valid:
                st.error(msg)
            else:
                st.success("✅ صيغة المفتاح صحيحة")
                st.session_state.gemini_key=key_input
        elif key_input != st.session_state.gemini_key:
            st.session_state.gemini_key=key_input

        c1,c2=st.columns(2)
        st.session_state.professor_name=c1.text_input(
            "👤 اسم الأستاذ المسؤول",value=st.session_state.professor_name,
            placeholder="أ.د / محمد أحمد...")
        st.session_state.report_date=c2.text_input(
            "📅 تاريخ الجرد",value=st.session_state.report_date)

    st.markdown("---")
    st.markdown('<div class="step-label">📸 الخطوة 1 — تحميل صور الجهاز</div>',
                unsafe_allow_html=True)

    c1,c2=st.columns(2)
    with c1:
        prim=st.file_uploader("⭐ الصورة الرئيسية (تظهر في التقرير)",
            type=["jpg","jpeg","png","webp"],key="up_primary")
    with c2:
        refs=st.file_uploader("📎 صور مرجعية (لوحة البيانات، أسفل الجهاز)",
            type=["jpg","jpeg","png","webp"],accept_multiple_files=True,key="up_refs")

    ca,cb=st.columns(2)
    with ca:
        if st.button("📥 تثبيت الصور"):
            added=0
            if prim: add_photo(resize_img(prim.read()),True,"الصورة الرئيسية"); added+=1
            if refs:
                for i,rf in enumerate(refs):
                    add_photo(resize_img(rf.read()),False,f"مرجعية {i+1}"); added+=1
            if added: st.success(f"✅ أُضيفت {added} صورة"); st.rerun()
    with cb:
        if st.button("🗑️ مسح جميع الصور"):
            st.session_state.captured_photos=[]
            st.session_state.photo_hashes=set()
            st.session_state.edit_photo_idx=None
            st.rerun()

    shots=st.session_state.captured_photos
    if shots:
        st.markdown(f"**{len(shots)} صورة مُثبَّتة:**")

        if st.session_state.edit_photo_idx is not None:
            idx=st.session_state.edit_photo_idx
            if 0<=idx<len(shots):
                image_editor(idx)
            else:
                st.session_state.edit_photo_idx=None
        else:
            cols=st.columns(min(len(shots),3))
            to_del=[]
            for i,ph in enumerate(shots):
                with cols[i%3]:
                    badge_cls="badge-primary" if ph["is_primary"] else "badge-ref"
                    st.markdown(f"""
                    <div class="photo-card">
                      <span class="{badge_cls}">{ph['label']}</span><br><br>
                      <img src="data:image/jpeg;base64,{b64img(ph['data'])}"
                           style="width:100%;border-radius:8px;max-height:140px;object-fit:cover;"/>
                    </div>""",unsafe_allow_html=True)
                    ea,eb=st.columns(2)
                    with ea:
                        if st.button("✏️ تعديل",key=f"ed_{ph['id']}"):
                            st.session_state.edit_photo_idx=i; st.rerun()
                    with eb:
                        if st.button("🗑️ حذف",key=f"dl_{ph['id']}"):
                            to_del.append(ph["id"])
            if to_del:
                st.session_state.photo_hashes={
                    hash(p["data"]) for p in shots if p["id"] not in to_del}
                st.session_state.captured_photos=[
                    p for p in shots if p["id"] not in to_del]
                st.rerun()

        # Quick AI fix all
        if st.session_state.gemini_key:
            if st.button("🤖 تصحيح تلقائي لجميع الصور (EXIF + تحسين)"):
                with st.spinner("⏳ معالجة الصور..."):
                    for i in range(len(st.session_state.captured_photos)):
                        d=st.session_state.captured_photos[i]["data"]
                        st.session_state.captured_photos[i]["data"]=process_image(d,enhance=True)
                st.success("✅ تمت معالجة جميع الصور"); st.rerun()

    st.markdown("---")
    st.markdown('<div class="step-label">🤖 الخطوة 2 — الاستخراج الذكي بـ Gemini AI</div>',
                unsafe_allow_html=True)

    ca,cb=st.columns(2)
    with ca:
        if st.button("🚀 استخراج البيانات بالذكاء الاصطناعي",
                     disabled=not shots,type="primary"):
            if not st.session_state.gemini_key:
                st.error("⚠️ أدخل Gemini API Key أولاً")
            else:
                valid,msg=validate_key(st.session_state.gemini_key)
                if not valid:
                    st.error(msg)
                else:
                    with st.spinner("🔍 Gemini يحلل الصور..."):
                        res=extract_device_info(
                            [p["data"] for p in shots],st.session_state.gemini_key)
                        st.session_state.ai_result=res
                    if "error" not in res: st.success("✅ تم الاستخراج!")
                    st.rerun()
    with cb:
        if st.button("✏️ إدخال يدوي"):
            st.session_state.ai_result={
                "device_type":"","brand":"","serial_number":"",
                "condition":"جيد","notes":"","confidence":"manual"}
            st.rerun()

    ai=st.session_state.ai_result or {}
    if "confidence" in ai and ai["confidence"]!="manual":
        cf={"high":("🟢","عالية"),"medium":("🟡","متوسطة"),"low":("🔴","منخفضة")}
        ic,lb=cf.get(ai["confidence"],("⚪","—"))
        st.info(f"{ic} دقة الاستخراج: **{lb}**")
    if "error" in ai:
        st.error(ai["error"])

    st.markdown("---")

    if ai and "device_type" in ai:
        st.markdown('<div class="step-label">📝 الخطوة 3 — مراجعة وتأكيد البيانات</div>',
                    unsafe_allow_html=True)
        fc1,fc2=st.columns(2)
        device_type=fc1.text_input("🔧 نوع الجهاز",value=ai.get("device_type",""))
        brand=fc2.text_input("🏭 الماركة / المصنّع",value=ai.get("brand",""))
        serial=fc1.text_input("🔢 الرقم التسلسلي",value=ai.get("serial_number",""))
        ci=STATUS_OPT.index(ai.get("condition","جيد")) if ai.get("condition","") in STATUS_OPT else 2
        condition=fc2.selectbox("📊 الحالة",STATUS_OPT,index=ci)
        notes=st.text_area("📒 ملاحظات الأستاذ",value=ai.get("notes",""),
                            height=95,placeholder="أي ملاحظات إضافية...")

        if is_duplicate(serial):
            st.markdown(f"""<div class="dup-warning">
              ⚠️ <strong>تحذير: الرقم التسلسلي مكرر!</strong><br>
              الرقم <code>{serial}</code> مسجّل مسبقاً — تحقق قبل الإضافة.
            </div>""",unsafe_allow_html=True)

        st.markdown("---")
        s1,s2=st.columns([3,1])
        with s1:
            if st.button("✅ إضافة إلى قائمة الجرد",type="primary"):
                st.session_state.inventory.append({
                    "id":uuid.uuid4().hex,"device_type":device_type,"brand":brand,
                    "serial_number":serial,"condition":condition,"notes":notes,
                    "photos":list(st.session_state.captured_photos),
                    "is_duplicate":is_duplicate(serial),
                    "added_at":datetime.datetime.now().strftime("%H:%M"),
                })
                for k in ["captured_photos","ai_result","edit_photo_idx"]:
                    st.session_state[k]=[] if k=="captured_photos" else None
                st.session_state.photo_hashes=set()
                st.markdown('<div class="success-flash">✅ تمت الإضافة بنجاح!</div>',
                            unsafe_allow_html=True)
                st.rerun()
        with s2:
            if st.button("↩️ تعيين"):
                for k in ["captured_photos","ai_result","edit_photo_idx"]:
                    st.session_state[k]=[] if k=="captured_photos" else None
                st.session_state.photo_hashes=set(); st.rerun()


# ================================================================
# 📋  PAGE: LIST
# ================================================================
def page_list():
    inv=st.session_state.inventory
    if not inv:
        st.info("📭 قائمة الجرد فارغة — أضف أجهزة من صفحة 'إضافة جهاز'.")
        return
    dups=len(get_duplicates())
    c1,c2,c3,c4=st.columns(4)
    c1.metric("📦 إجمالي الأجهزة",len(inv))
    c2.metric("🏭 الماركات",len(set(i.get("brand","") for i in inv)))
    c3.metric("⚠️ تحتاج صيانة",sum(1 for i in inv if i.get("condition") in ["يحتاج صيانة","معطل"]))
    c4.metric("🔄 مكررة",dups)
    if dups: st.error(f"⚠️ يوجد **{dups}** رقم تسلسلي مكرر")
    st.markdown("---")
    for i,item in enumerate(inv):
        dup=item.get("is_duplicate") or is_duplicate(item.get("serial_number",""),i)
        ph=next((p["data"] for p in item.get("photos",[]) if p.get("is_primary")),None)
        if ph is None and item.get("photos"): ph=item["photos"][0]["data"]
        ci,cd,cx=st.columns([1,4,1])
        with ci:
            if ph: st.image(thumb(ph,(120,120)),use_container_width=True)
        with cd:
            dup_b="  🔴 **مكرر**" if dup else ""
            st.markdown(f"**{i+1}. {item.get('device_type','غير محدد')}**{dup_b}")
            st.markdown(f"🏭 `{item.get('brand','—')}` | 🔢 `{item.get('serial_number','—')}` | 📊 {item.get('condition','—')} | 🕐 {item.get('added_at','')}")
            if item.get("notes"): st.caption(f"📝 {item['notes'][:100]}")
        with cx:
            if st.button("🗑️",key=f"di_{i}"):
                st.session_state.inventory.pop(i); st.rerun()
        st.markdown("---")


# ================================================================
# 📤  PAGE: EXPORT
# ================================================================
def page_export():
    inv=st.session_state.inventory
    if not inv:
        st.warning("⚠️ أضف أجهزة إلى قائمة الجرد أولاً."); return

    dups=get_duplicates()
    st.markdown(f"""
    <div style="background:#eaf0fb;border-radius:12px;padding:1.1rem;direction:rtl;
                border-right:4px solid #1a5276;margin-bottom:1rem;">
      <strong>👤</strong> المُعِد: {st.session_state.professor_name or '<em>غير محدد</em>'}<br>
      <strong>📅</strong> التاريخ: {st.session_state.report_date}<br>
      <strong>📦</strong> عدد الأجهزة: {len(inv)} &nbsp;|&nbsp;
      <strong>⚠️</strong> أرقام مكررة: {len(dups)}
    </div>""",unsafe_allow_html=True)

    with st.spinner("⏳ تحميل الخطوط العربية..."):
        fonts_ok,fonts_msg=setup_fonts()

    if not fonts_ok:
        st.error(f"❌ فشل تحميل خطوط PDF: {fonts_msg}")
        st.warning("💡 تأكد من الاتصال بالإنترنت — خط Cairo يُحمَّل تلقائياً عند أول استخدام.")
    else:
        st.success("✅ الخطوط العربية جاهزة")

    st.markdown("---")
    cp,cd=st.columns(2)
    with cp:
        st.markdown("#### 📄 ملف PDF")
        st.caption("للأرشفة الرسمية والطباعة")
        if st.button("🖨️ توليد PDF",disabled=not fonts_ok,type="primary"):
            with st.spinner("⏳ جاري إنشاء PDF..."):
                try:
                    pdf=PDFReport(inv,st.session_state.professor_name,
                                   st.session_state.report_date).build()
                    fn=f"جرد_{LAB_NAME}_{st.session_state.report_date.replace('/','-')}.pdf"
                    st.download_button("⬇️ تحميل PDF",pdf,fn,"application/pdf",key="dl_pdf")
                    st.success("✅ PDF جاهز!")
                except Exception as e:
                    st.error(f"❌ خطأ PDF: {e}")
                    import traceback; st.code(traceback.format_exc(),language="text")
    with cd:
        st.markdown("#### 📝 ملف Word (DOCX)")
        st.caption("للتعديل والمراجعة")
        if st.button("📄 توليد DOCX"):
            with st.spinner("⏳ جاري إنشاء DOCX..."):
                try:
                    docx_b=DOCXReport(inv,st.session_state.professor_name,
                                       st.session_state.report_date).build()
                    fn=f"جرد_{LAB_NAME}_{st.session_state.report_date.replace('/','-')}.docx"
                    MIME="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    st.download_button("⬇️ تحميل DOCX",docx_b,fn,MIME,key="dl_docx")
                    st.success("✅ DOCX جاهز!")
                except Exception as e:
                    st.error(f"❌ خطأ DOCX: {e}")
                    import traceback; st.code(traceback.format_exc(),language="text")

    st.markdown("---")
    with st.expander("⚠️ خيارات متقدمة"):
        if st.button("🗑️ مسح قائمة الجرد بالكامل"):
            for k in ["inventory","captured_photos","photo_hashes","ai_result"]:
                st.session_state[k]=[] if k!="photo_hashes" else set()
            st.success("✅ تم المسح"); st.rerun()


# ================================================================
# 🚀  MAIN
# ================================================================
def main():
    st.markdown(f"""
    <div class="app-header">
      <h1>🔭 نظام جرد معمل المساحة</h1>
      <p>{UNIV_NAME} | {FAC_NAME} | مدعوم بـ Gemini AI</p>
    </div>""",unsafe_allow_html=True)
    nav()
    st.markdown("")
    pg=st.session_state.page
    if   pg=="capture": page_capture()
    elif pg=="list":    page_list()
    elif pg=="export":  page_export()

if __name__=="__main__":
    main()
